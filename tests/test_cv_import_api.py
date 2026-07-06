"""POST /api/cv/import: multipart PDF/DOCX → reviewable master_cv.draft.yaml (F2 wizard).

Deterministic ($0) text extraction only — reuses engine/cv/import_cv.py. The endpoint writes
ONLY master_cv.draft.yaml (never the real master_cv.yaml), is origin-guarded, and rejects
unsupported/empty/corrupt uploads with a 4xx (never a 500). All fixtures are 100% fictitious.
"""

from __future__ import annotations

from fastapi.testclient import TestClient

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _redirect_profile_dir(monkeypatch, tmp_path):
    """Point MASTER_CV_PATH at a throwaway profile dir so the draft never lands in the repo."""
    import engine.paths as paths

    monkeypatch.setattr(paths, "MASTER_CV_PATH", tmp_path / "profile" / "master_cv.yaml")


def _docx_bytes(tmp_path) -> bytes:
    from docx import Document

    p = tmp_path / "cv.docx"
    d = Document()
    d.add_paragraph("Test Candidate — Data Engineer")
    d.add_paragraph("Experience: built pipelines at FicticiaCorp.")
    d.save(str(p))
    return p.read_bytes()


def _pdf_bytes(tmp_path) -> bytes:
    from reportlab.pdfgen import canvas

    p = tmp_path / "cv.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(72, 720, "Test Candidate — Data Engineer")
    c.drawString(72, 700, "Experience: built pipelines at FicticiaCorp.")
    c.save()
    return p.read_bytes()


def test_import_docx_returns_draft_and_writes_only_the_draft(atlas_app, tmp_path, monkeypatch):
    _redirect_profile_dir(monkeypatch, tmp_path)
    with TestClient(atlas_app) as client:
        resp = client.post(
            "/api/cv/import",
            files={"file": ("cv.docx", _docx_bytes(tmp_path), _DOCX_MIME)},
        )
    assert resp.status_code == 200
    body = resp.json()
    assert body["ok"] is True and body["chars"] > 0
    assert "Test Candidate" in body["draft"]  # raw text landed in _source_text
    assert "_source_text" in body["draft"]
    draft_path = tmp_path / "profile" / "master_cv.draft.yaml"
    assert draft_path.exists()
    assert body["path"] == str(draft_path)
    # Safety invariant: NEVER writes the real CV.
    assert not (tmp_path / "profile" / "master_cv.yaml").exists()


def test_import_pdf_returns_draft(atlas_app, tmp_path, monkeypatch):
    _redirect_profile_dir(monkeypatch, tmp_path)
    with TestClient(atlas_app) as client:
        resp = client.post(
            "/api/cv/import",
            files={"file": ("cv.pdf", _pdf_bytes(tmp_path), "application/pdf")},
        )
    assert resp.status_code == 200
    body = resp.json()
    assert body["ok"] is True and body["chars"] > 0
    assert "Test Candidate" in body["draft"]
    assert (tmp_path / "profile" / "master_cv.draft.yaml").exists()
    assert not (tmp_path / "profile" / "master_cv.yaml").exists()


def test_import_rejects_unsupported_format(atlas_app, tmp_path, monkeypatch):
    _redirect_profile_dir(monkeypatch, tmp_path)
    with TestClient(atlas_app) as client:
        resp = client.post(
            "/api/cv/import", files={"file": ("cv.txt", b"plain text", "text/plain")}
        )
    assert resp.status_code == 400
    # Nothing written for a rejected format.
    assert not (tmp_path / "profile" / "master_cv.draft.yaml").exists()


def test_import_corrupt_file_is_graceful(atlas_app, tmp_path, monkeypatch):
    """A .docx/.pdf that is not a real document must 4xx gracefully, never 500."""
    _redirect_profile_dir(monkeypatch, tmp_path)
    with TestClient(atlas_app) as client:
        resp = client.post(
            "/api/cv/import",
            files={"file": ("cv.docx", b"not a real docx at all", _DOCX_MIME)},
        )
    assert resp.status_code == 400
    assert not (tmp_path / "profile" / "master_cv.draft.yaml").exists()


def test_import_empty_extraction_is_graceful(atlas_app, tmp_path, monkeypatch):
    """A structurally valid but text-empty document (e.g. a blank/scanned PDF) → 400, not 500."""
    from reportlab.pdfgen import canvas

    _redirect_profile_dir(monkeypatch, tmp_path)
    p = tmp_path / "blank.pdf"
    canvas.Canvas(str(p)).save()  # a valid PDF with no text
    with TestClient(atlas_app) as client:
        resp = client.post(
            "/api/cv/import",
            files={"file": ("blank.pdf", p.read_bytes(), "application/pdf")},
        )
    assert resp.status_code == 400
    assert not (tmp_path / "profile" / "master_cv.draft.yaml").exists()


def test_import_rejects_foreign_origin(atlas_app):
    with TestClient(atlas_app) as client:
        resp = client.post(
            "/api/cv/import",
            files={"file": ("cv.docx", b"x", "application/octet-stream")},
            headers={"origin": "https://evil.example.com"},
        )
    assert resp.status_code == 403
