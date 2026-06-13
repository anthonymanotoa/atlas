"""FASE C: import an existing CV (PDF/DOCX) → deterministic text → reviewable draft."""

from __future__ import annotations

import pytest
import yaml

from engine.cv.import_cv import build_draft, extract_text


def test_extract_text_from_docx(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Ada Lovelace — Senior Data Scientist")
    doc.add_paragraph("Skills: Python, SQL, Machine Learning")
    p = tmp_path / "cv.docx"
    doc.save(str(p))
    text = extract_text(p)
    assert "Ada Lovelace" in text
    assert "Machine Learning" in text


def test_extract_text_from_pdf(tmp_path):
    from reportlab.pdfgen import canvas

    p = tmp_path / "cv.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(72, 720, "Grace Hopper — Compiler Pioneer")
    c.drawString(72, 700, "Skills: COBOL, Systems")
    c.save()
    text = extract_text(p)
    assert "Grace Hopper" in text


def test_unsupported_format_raises(tmp_path):
    p = tmp_path / "cv.txt"
    p.write_text("plain text resume")
    with pytest.raises(ValueError, match="unsupported CV format"):
        extract_text(p)


def test_build_draft_embeds_source_and_scaffold():
    raw = "Ada Lovelace\nPython, SQL"
    draft = build_draft(raw)
    assert draft.startswith("# DRAFT")  # explicit not-your-master_cv banner
    loaded = yaml.safe_load(draft)
    assert loaded["_source_text"] == raw  # raw extraction preserved for Cowork to map
    # scaffold shape matches the master_cv schema, but is empty (no fabrication)
    assert set(loaded["basics"]) >= {"name", "label", "email", "summary"}
    assert loaded["skills"] == [] and loaded["experience"] == []
