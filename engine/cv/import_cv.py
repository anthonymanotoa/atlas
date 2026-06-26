"""Import an existing CV (PDF/DOCX) into a reviewable master_cv *draft*.

Two-step by design, and deliberately conservative:
  1. **Deterministic** (here): extract the plain text from the file. No structuring, no LLM,
     no guessing — CV layouts vary too much for a reliable deterministic parser.
  2. **Human + Claude Cowork**: map the extracted text into the `master_cv.yaml` schema
     (truthfully, never inventing), review it, and save it themselves.

This module therefore only ever writes a ``master_cv.draft.yaml`` and NEVER touches the real
``master_cv.yaml`` — the user owns that final step. Keeps the $0 model intact (Cowork is the
subscription client; nothing here calls an API).
"""

from __future__ import annotations

from pathlib import Path

import yaml

SUPPORTED = {".pdf", ".docx"}

# Mirrors profile/master_cv.example.yaml so Cowork has the exact target shape to fill.
_SCAFFOLD: dict = {
    "basics": {
        "name": "",
        "label": "",
        "email": "",
        "phone": "",
        "location": "",
        "linkedin": "",
        "github": "",
        "website": "",
        "summary": "",
    },
    "skills": [],
    "experience": [],  # items: {title, company, location, start, end, skills: [], highlights: []}
    "education": [],  # items: {degree, area, institution, start, end}
    "certifications": [],  # items: {name, issuer, date}
    "projects": [],  # items: {name, description, highlights: []}
}

_HEADER = (
    "# DRAFT — imported CV (NOT your master_cv.yaml).\n"
    "# The engine only extracted raw text into `_source_text` below. Ask Claude (Cowork) to map\n"
    "# it into the fields above TRUTHFULLY (never invent experience), review it yourself, then\n"
    "# save the result as profile/master_cv.yaml. Delete `_source_text` once mapped.\n"
)


def extract_text(path: Path) -> str:
    """Deterministically extract plain text from a PDF or DOCX. No structuring, no LLM."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    raise ValueError(f"unsupported CV format {ext!r}; supported: {sorted(SUPPORTED)}")


def _extract_pdf(path: Path) -> str:
    import pdfplumber  # lazy: optional dep, only needed for PDF import

    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts).strip()


def _extract_docx(path: Path) -> str:
    from docx import Document  # python-docx (already a dependency)

    doc = Document(str(path))
    lines = [p.text for p in doc.paragraphs]
    # Tables (skills grids, two-column layouts) commonly hold CV content too.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(line for line in lines if line.strip()).strip()


def _scaffold_for(domain: str | None) -> dict:
    """The empty master_cv scaffold for a domain. Architecture surfaces the extra fields a
    regulated, portfolio-driven profile needs (pitch, portfolio link, licensure/registration)
    so the Cowork mapper knows to fill them. Other domains get the neutral base scaffold."""
    sc: dict = {
        "basics": dict(_SCAFFOLD["basics"]),
        "skills": [],
        "experience": [],
        "education": [],
        "certifications": [],
        "projects": [],
    }
    if domain == "architecture":
        sc["basics"]["portfolio"] = ""  # near-mandatory in architecture
        sc["basics"]["pitch"] = {
            "identity_line": "",
            "role_noun": "",
            "impact_domain": "",
            "value_verb": "",
        }
        # items: {title, issuer, status} — e.g. título de Arquitecta, registro SENESCYT/CAE
        sc["licensure"] = []
    return sc


def build_draft(source_text: str, domain: str | None = None) -> str:
    """A YAML draft: the (per-domain) master_cv scaffold + the raw extracted text, for Cowork+human
    to map. ``domain`` selects which fields the scaffold surfaces (defaults to the neutral base)."""
    doc = {**_scaffold_for(domain), "_source_text": source_text}
    return _HEADER + yaml.safe_dump(doc, allow_unicode=True, sort_keys=False)
