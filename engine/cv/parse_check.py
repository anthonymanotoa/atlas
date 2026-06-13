"""Re-parse the generated DOCX to confirm an ATS would read it correctly.

Emulates what Lever/Workday store: extract plain text + structure and assert the
contact info, standard section headings and dates survived. Cheap local equivalent
of an ATS-screener second opinion.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from engine.cv.render import HEADINGS

_DATE = re.compile(
    r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|\b(19|20)\d{2}\b", re.I
)
_EMAIL = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")


def extract_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    return "\n".join(p.text for p in doc.paragraphs)


def check(docx_path: Path, cv: dict, language: str = "en") -> tuple[bool, list[str]]:
    doc = Document(str(docx_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    low = text.lower()
    issues: list[str] = []

    name = (cv.get("basics", {}) or {}).get("name", "")
    if name and name.lower() not in low:
        issues.append("name not found in parsed text")
    if not _EMAIL.search(text):
        issues.append("no email parsed (check it's in the body, not a header)")

    h = HEADINGS.get(language, HEADINGS["en"])
    for key in ("summary", "skills", "experience"):
        if h[key].lower() not in low:
            issues.append(f"missing section heading: {h[key]}")
    if not _DATE.search(text):
        issues.append("no parseable dates (use 'Mon YYYY')")
    if doc.tables:
        issues.append(f"{len(doc.tables)} table(s) present — ATS parsers garble tables")

    return (len(issues) == 0), issues
