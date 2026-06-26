"""Render a CV dict to a parse-safe DOCX (and optional PDF).

Parse-safety rules baked in (from ATS-parser research): single column, standard
section headings, Month-YYYY dates, no tables/text-boxes/headers-footers/graphics,
contact info in the body, reverse-chronological. DOCX is the default (best parse rate).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from engine.config import load_cv_layout

DEFAULT_ORDER = ["summary", "skills", "experience", "education", "certs", "projects"]

HEADINGS = {
    "en": {
        "summary": "Professional Summary",
        "skills": "Skills",
        "experience": "Experience",
        "education": "Education",
        "certs": "Certifications",
        "projects": "Projects",
        "licensure": "Licensure & Registration",
    },
    "es": {
        "summary": "Resumen Profesional",
        "skills": "Habilidades",
        "experience": "Experiencia",
        "education": "Educación",
        "certs": "Certificaciones",
        "projects": "Proyectos",
        "licensure": "Licenciatura y Registro",
    },
}


def _heading(key: str, language: str, layout: dict) -> str:
    """Section heading for `key`, honoring per-profile cv_layout.yaml label overrides."""
    base = HEADINGS.get(language, HEADINGS["en"]).get(key, key.title())
    override = (layout.get("labels") or {}).get(key) or {}
    return override.get(language) or override.get("en") or base


def _licensure_line(it: dict) -> str:
    """One Licensure/Registration entry → 'Title — Issuer — Status' (architecture, etc.)."""
    return " — ".join(
        p
        for p in [it.get("title") or it.get("name"), it.get("issuer"), it.get("status") or it.get("date")]
        if p
    )


def _section(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def _contact_line(basics: dict) -> str:
    parts = [
        basics.get("email"),
        basics.get("phone"),
        basics.get("location"),
        basics.get("linkedin"),
        basics.get("github"),
        basics.get("website"),
    ]
    return "  |  ".join(p for p in parts if p)


def render_docx(cv: dict, out_path: Path, language: str = "en", layout: dict | None = None) -> Path:
    layout = layout or load_cv_layout()
    order = layout.get("order") or DEFAULT_ORDER
    basics = cv.get("basics", {}) or {}
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(40)
        section.left_margin = section.right_margin = Pt(54)

    # Name + target title + contact (all in the body, no header/footer).
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name_p.add_run(basics.get("name", ""))
    r.bold = True
    r.font.size = Pt(20)
    if basics.get("label"):
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(basics["label"])
        lr.italic = True
        lr.font.size = Pt(12)
    cp = doc.add_paragraph(_contact_line(basics))
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def head(key: str) -> None:
        _section(doc, _heading(key, language, layout))

    def summary() -> None:
        if basics.get("summary"):
            head("summary")
            doc.add_paragraph(" ".join(basics["summary"].split()))

    def skills() -> None:
        items = cv.get("skills") or []
        if items:
            head("skills")
            doc.add_paragraph("  ·  ".join(items))  # single inline line — no tables/columns

    def experience() -> None:
        exp = cv.get("experience") or []
        if not exp:
            return
        head("experience")
        for e in exp:
            hp = doc.add_paragraph()
            hr = hp.add_run(f"{e.get('title', '')} — {e.get('company', '')}")
            hr.bold = True
            meta = "  |  ".join(
                p
                for p in [e.get("location"), f"{e.get('start', '')} – {e.get('end', '')}".strip(" –")]
                if p
            )
            if meta:
                mr = hp.add_run(f"\n{meta}")
                mr.italic = True
                mr.font.size = Pt(10)
            for hl in e.get("highlights") or []:
                doc.add_paragraph(" ".join(hl.split()), style="List Bullet")

    def education() -> None:
        edu = cv.get("education") or []
        if not edu:
            return
        head("education")
        for ed in edu:
            line = " ".join(
                p
                for p in [
                    ed.get("degree"),
                    ed.get("area"),
                    "—" if ed.get("institution") else None,
                    ed.get("institution"),
                ]
                if p
            )
            dates = f"{ed.get('start', '')} – {ed.get('end', '')}".strip(" –")
            p = doc.add_paragraph()
            p.add_run(line).bold = True
            if dates:
                dr = p.add_run(f"  ({dates})")
                dr.font.size = Pt(10)

    def certs() -> None:
        rows = [c for c in (cv.get("certifications") or []) if c.get("name")]
        if not rows:
            return
        head("certs")
        for c in rows:
            txt = " — ".join(p for p in [c.get("name"), c.get("issuer"), c.get("date")] if p)
            doc.add_paragraph(txt, style="List Bullet")

    def licensure() -> None:
        rows = [it for it in (cv.get("licensure") or []) if it.get("title") or it.get("name")]
        if not rows:
            return
        head("licensure")
        for it in rows:
            doc.add_paragraph(_licensure_line(it), style="List Bullet")

    def projects() -> None:
        rows = cv.get("projects") or []
        if not rows:
            return
        head("projects")
        for pr in rows:
            p = doc.add_paragraph()
            p.add_run(pr.get("name", "")).bold = True
            if pr.get("description"):
                doc.add_paragraph(" ".join(pr["description"].split()))
            for hl in pr.get("highlights") or []:
                doc.add_paragraph(" ".join(hl.split()), style="List Bullet")

    renderers = {
        "summary": summary, "skills": skills, "experience": experience, "education": education,
        "certs": certs, "licensure": licensure, "projects": projects,
    }
    for key in order:
        fn = renderers.get(key)
        if fn:
            fn()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _esc(s: str) -> str:
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def render_pdf(
    cv: dict, out_path: Path, language: str = "en", layout: dict | None = None
) -> Path | None:
    """Render a clean single-column PDF (reportlab — no LibreOffice/Word needed)."""
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate

    layout = layout or load_cv_layout()
    order = layout.get("order") or DEFAULT_ORDER
    basics = cv.get("basics", {}) or {}
    # NOTE: reportlab's ParagraphStyle defaults `leading` to a FIXED 12pt when omitted (it is
    # NOT auto-scaled to the font size). So any style with fontSize > ~11 MUST set `leading`
    # explicitly, or its line box is shorter than the glyphs and the next paragraph is drawn
    # on top of it — which is exactly what made the name and the title line overlap. Every
    # style below sets leading ≈ 1.2–1.3 × fontSize.
    name_s = ParagraphStyle(
        "name",
        fontName="Helvetica-Bold",
        fontSize=19,
        leading=23,
        alignment=TA_CENTER,
        spaceAfter=3,
    )
    label_s = ParagraphStyle(
        "label",
        fontName="Helvetica-Oblique",
        fontSize=11,
        leading=14,
        alignment=TA_CENTER,
        textColor="#333333",
        spaceAfter=2,
    )
    contact_s = ParagraphStyle(
        "contact",
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        alignment=TA_CENTER,
        textColor="#444444",
        spaceAfter=8,
    )
    head_s = ParagraphStyle(
        "head",
        fontName="Helvetica-Bold",
        fontSize=11.5,
        leading=15,
        textColor="#1a1a1a",
        spaceBefore=10,
        spaceAfter=3,
    )
    body_s = ParagraphStyle("body", fontName="Helvetica", fontSize=10, leading=13)
    role_s = ParagraphStyle(
        "role", fontName="Helvetica-Bold", fontSize=10.5, leading=13, spaceBefore=5
    )
    meta_s = ParagraphStyle(
        "meta",
        fontName="Helvetica-Oblique",
        fontSize=9,
        leading=12,
        textColor="#555555",
        spaceAfter=2,
    )
    bullet_s = ParagraphStyle("bullet", fontName="Helvetica", fontSize=9.5, leading=12.5)

    story = [Paragraph(_esc(basics.get("name", "")), name_s)]
    if basics.get("label"):
        story.append(Paragraph(_esc(basics["label"]), label_s))
    story.append(Paragraph(_esc(_contact_line(basics)), contact_s))

    def bullets(items):
        return ListFlowable(
            [
                ListItem(Paragraph(_esc(" ".join(i.split())), bullet_s), leftIndent=12)
                for i in items
            ],
            bulletType="bullet",
            start="•",
            leftIndent=10,
        )

    def hd(key: str):
        return Paragraph(_heading(key, language, layout).upper(), head_s)

    def summary():
        if basics.get("summary"):
            story.extend([hd("summary"), Paragraph(_esc(" ".join(basics["summary"].split())), body_s)])

    def skills():
        if cv.get("skills"):
            story.extend([hd("skills"), Paragraph("  ·  ".join(_esc(s) for s in cv["skills"]), body_s)])

    def experience():
        if not cv.get("experience"):
            return
        story.append(hd("experience"))
        for e in cv["experience"]:
            story.append(
                Paragraph(f"{_esc(e.get('title', ''))} — {_esc(e.get('company', ''))}", role_s)
            )
            meta = "  |  ".join(
                p
                for p in [
                    _esc(e.get("location")),
                    f"{e.get('start', '')} – {e.get('end', '')}".strip(" –"),
                ]
                if p
            )
            if meta:
                story.append(Paragraph(meta, meta_s))
            if e.get("highlights"):
                story.append(bullets(e["highlights"]))

    def education():
        edu = cv.get("education") or []
        if not edu:
            return
        story.append(hd("education"))
        for ed in edu:
            line = " ".join(
                p
                for p in [
                    ed.get("degree"),
                    ed.get("area"),
                    "—" if ed.get("institution") else None,
                    ed.get("institution"),
                ]
                if p
            )
            dates = f"{ed.get('start', '')} – {ed.get('end', '')}".strip(" –")
            story.append(
                Paragraph(f"<b>{_esc(line)}</b>" + (f"  ({dates})" if dates else ""), body_s)
            )

    def certs():
        rows = [c for c in (cv.get("certifications") or []) if c.get("name")]
        if not rows:
            return
        story.append(hd("certs"))
        story.append(
            bullets(
                [" — ".join(p for p in [c.get("name"), c.get("issuer"), c.get("date")] if p) for c in rows]
            )
        )

    def licensure():
        rows = [it for it in (cv.get("licensure") or []) if it.get("title") or it.get("name")]
        if not rows:
            return
        story.append(hd("licensure"))
        story.append(bullets([_licensure_line(it) for it in rows]))

    def projects():
        rows = cv.get("projects") or []
        if not rows:
            return
        story.append(hd("projects"))
        for pr in rows:
            story.append(Paragraph(f"<b>{_esc(pr.get('name', ''))}</b>", body_s))
            if pr.get("description"):  # parity with render_docx — PDF previously dropped this
                story.append(Paragraph(_esc(" ".join(pr["description"].split())), body_s))
            if pr.get("highlights"):
                story.append(bullets(pr["highlights"]))

    renderers = {
        "summary": summary, "skills": skills, "experience": experience, "education": education,
        "certs": certs, "licensure": licensure, "projects": projects,
    }
    for key in order:
        fn = renderers.get(key)
        if fn:
            fn()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
    )
    try:
        doc.build(story)
    except Exception as exc:  # noqa: BLE001 — never crash a prep run on a PDF render hiccup,
        # but log WHY so a silently-missing PDF is debuggable instead of a mystery.
        import logging

        logging.getLogger("atlas.cv").warning("PDF render failed for %s: %s", out_path.name, exc)
        return None
    return out_path if out_path.exists() else None
